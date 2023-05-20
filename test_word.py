from docx import Document

document = Document()

document.add_heading('ezpg', 0)
document.add_paragraph('csgo')

document.add_picture("a.png")

num = 0
for para in document.paragraphs:
        num = num + 1
        print(num, para.text)
print(num)

doc =document("sample.docx")
print(len(doc))
document.save('sample_answer.docx')